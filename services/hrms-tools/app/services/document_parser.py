"""
Document Parser Service
Extracts text from PDF and DOC/DOCX files
"""
import io
import logging
import base64
from typing import Tuple, Optional
import PyPDF2
import docx
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentParserError(Exception):
    """Exception raised when document parsing fails"""
    pass


class DocumentParser:
    """Handles parsing of CV documents in various formats"""

    @staticmethod
    def parse_base64(base64_content: str, filename: str) -> Tuple[str, int]:
        """
        Parse document from base64 encoded content

        Args:
            base64_content: Base64 encoded document
            filename: Original filename to determine type

        Returns:
            Tuple of (extracted_text, page_count)

        Raises:
            DocumentParserError: If parsing fails
        """
        try:
            # Decode base64
            file_bytes = base64.b64decode(base64_content)
            file_stream = io.BytesIO(file_bytes)

            # Determine file type from extension
            extension = Path(filename).suffix.lower()

            if extension == '.pdf':
                return DocumentParser._parse_pdf(file_stream)
            elif extension in ['.doc', '.docx']:
                return DocumentParser._parse_docx(file_stream)
            else:
                raise DocumentParserError(f"Unsupported file type: {extension}")

        except base64.binascii.Error as e:
            logger.error(f"Base64 decode error: {e}")
            raise DocumentParserError(f"Invalid base64 content: {str(e)}")
        except Exception as e:
            logger.error(f"Document parsing error: {e}")
            raise DocumentParserError(f"Failed to parse document: {str(e)}")

    @staticmethod
    def _parse_pdf(file_stream: io.BytesIO) -> Tuple[str, int]:
        """
        Extract text from PDF file

        Args:
            file_stream: Binary stream of PDF file

        Returns:
            Tuple of (extracted_text, page_count)
        """
        try:
            pdf_reader = PyPDF2.PdfReader(file_stream)
            page_count = len(pdf_reader.pages)

            text_content = []
            for page_num in range(page_count):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_content.append(text)

            full_text = "\n\n".join(text_content)

            if not full_text.strip():
                raise DocumentParserError("No text could be extracted from PDF. It may be scanned or image-based.")

            logger.info(f"Successfully extracted text from PDF ({page_count} pages, {len(full_text)} chars)")
            return full_text, page_count

        except PyPDF2.errors.PdfReadError as e:
            logger.error(f"PDF read error: {e}")
            raise DocumentParserError(f"Invalid or corrupted PDF file: {str(e)}")
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            raise DocumentParserError(f"Failed to parse PDF: {str(e)}")

    @staticmethod
    def _parse_docx(file_stream: io.BytesIO) -> Tuple[str, int]:
        """
        Extract text from DOCX file

        Args:
            file_stream: Binary stream of DOCX file

        Returns:
            Tuple of (extracted_text, page_count)
            Note: Page count is estimated for DOCX files
        """
        try:
            doc = docx.Document(file_stream)

            # Extract all paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(" | ".join(row_text))

            full_text = "\n\n".join(text_content)

            if not full_text.strip():
                raise DocumentParserError("No text could be extracted from DOCX file.")

            # Estimate page count (rough approximation: 500 words per page)
            word_count = len(full_text.split())
            estimated_pages = max(1, word_count // 500)

            logger.info(f"Successfully extracted text from DOCX (~{estimated_pages} pages, {len(full_text)} chars)")
            return full_text, estimated_pages

        except docx.opc.exceptions.PackageNotFoundError as e:
            logger.error(f"DOCX format error: {e}")
            raise DocumentParserError("Invalid DOCX file format")
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            raise DocumentParserError(f"Failed to parse DOCX: {str(e)}")

    @staticmethod
    def validate_file_size(base64_content: str, max_size_mb: int = 10) -> bool:
        """
        Validate file size

        Args:
            base64_content: Base64 encoded file
            max_size_mb: Maximum allowed size in MB

        Returns:
            True if valid, False otherwise
        """
        try:
            # Decode to get actual file size
            file_bytes = base64.b64decode(base64_content)
            size_mb = len(file_bytes) / (1024 * 1024)

            if size_mb > max_size_mb:
                logger.warning(f"File size ({size_mb:.2f}MB) exceeds limit ({max_size_mb}MB)")
                return False

            return True
        except Exception as e:
            logger.error(f"File size validation error: {e}")
            return False

    @staticmethod
    def get_text_preview(text: str, max_chars: int = 500) -> str:
        """
        Get a preview of extracted text

        Args:
            text: Full text
            max_chars: Maximum characters in preview

        Returns:
            Text preview
        """
        if len(text) <= max_chars:
            return text

        return text[:max_chars] + "..."


# Singleton instance
_parser_instance: Optional[DocumentParser] = None


def get_document_parser() -> DocumentParser:
    """Get or create the global document parser instance"""
    global _parser_instance
    if _parser_instance is None:
        _parser_instance = DocumentParser()
    return _parser_instance
